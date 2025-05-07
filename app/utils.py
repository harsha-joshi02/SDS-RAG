from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import time
import pdfplumber
import logging
import os
from app.config import CONFIG

logger = logging.getLogger(__name__)

def load_sds(file_path: str):
    """
    Loads and extracts text (including tables) from a PDF or DOCX file.

    Supports:
    - PDFs: Extracts text and tables from all pages.
    - DOCX: Extracts paragraphs and table contents.

    Args:
        file_path (str): Path to the document file.

    Returns:
        str: Combined extracted text.

    Raises:
        ValueError: If the file format is unsupported.
        Exception: If any error occurs during processing.
    """

    start_time = time.time()
    text = ""

    try:
        if file_path.lower().endswith('.pdf'):
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
                    logger.info(f"Extracted {len(page_text)} characters from PDF page")

                    tables = page.extract_tables()
                    if tables:
                        logger.info(f"Extracted {len(tables)} tables from PDF page")
                    for table in tables:
                        for row in table:
                            text += " | ".join(str(cell).strip() if cell else "" for cell in row) + "\n"

        elif file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text.strip() + "\n"

            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() if cell.text.strip() else " " for cell in row.cells)
                    text += row_text + "\n"
            logger.info(f"Extracted {table_count} tables from DOCX")

        else:
            raise ValueError(f"Unsupported file type: {file_path}")
    except Exception as e:
        logger.error(f"Error processing {file_path}: {str(e)}", exc_info=True)
        raise

    end_time = time.time()
    logger.info(f"Document Processing Time: {end_time - start_time:.2f} sec")
    logger.info(f"Extracted {len(text)} characters from {os.path.basename(file_path)}")
    return text

def preprocess_text(text: str):
    """
    Splits input text into smaller chunks using recursive character-based splitting.

    Args:
        text (str): The input text to be chunked.

    Returns:
        List[str]: A list of text chunks. If text is too short, returns a single-element list.
    """

    if not text or len(text.strip()) < 100:
        logger.warning("Very short or empty text provided for chunking")
        return [text.strip()] if text.strip() else []
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CONFIG["text_splitter"]["chunk_size"], 
        chunk_overlap=CONFIG["text_splitter"]["chunk_overlap"],
    )
    chunks = text_splitter.split_text(text)
    
    logger.info(f"Created {len(chunks)} chunks from {len(text)} characters of text")
    
    for i, chunk in enumerate(chunks[:3]):
        logger.info(f"Chunk {i+1} preview: {chunk[:50]}...")
    
    if not chunks and text.strip():
        logger.warning("Chunking failed, using full text as a single chunk")
        chunks = [text.strip()]
    
    return chunks    
